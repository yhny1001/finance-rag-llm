"""
RAG引擎核心实现
包含文档处理、向量检索、答案生成等功能
集成FAISS向量数据库实现高效持久化
"""

import os
import json
import gc
import torch
from pathlib import Path
from typing import List, Dict, Any, Optional
from tqdm import tqdm

# 文档处理依赖
try:
    from docx import Document
    DOCX_AVAILABLE = True
    print("✅ python-docx可用")
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx不可用，.docx文件将无法正确处理")

try:
    import PyPDF2
    PDF_AVAILABLE = True
    print("✅ PyPDF2可用")
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️ PyPDF2不可用，.pdf文件将无法正确处理")

# 简化导入，避免llama-index的循环导入问题
try:
    # 直接导入transformers和sentence-transformers作为替代
    from transformers import AutoTokenizer, AutoModelForCausalLM
    from sentence_transformers import SentenceTransformer
    TRANSFORMERS_AVAILABLE = True
    print("✅ transformers和sentence-transformers导入成功")
except ImportError as e:
    print(f"⚠️ transformers导入失败: {e}")
    TRANSFORMERS_AVAILABLE = False

# 尝试导入jieba
try:
    import jieba
    JIEBA_AVAILABLE = True
    print("✅ jieba导入成功")
except ImportError:
    print("⚠️ jieba未安装，文本分词功能将受限")
    JIEBA_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
    print("✅ pandas导入成功")
except ImportError:
    print("⚠️ pandas未安装")
    PANDAS_AVAILABLE = False

from config import Config
from vector_db import VectorDatabase


class SimpleDocument:
    """简化的文档类，替代llama-index的Document"""
    def __init__(self, text: str, metadata: Dict[str, Any] = None):
        self.text = text
        self.metadata = metadata or {}


class DocumentProcessor:
    """文档处理器"""
    
    def __init__(self):
        self.supported_formats = ['.txt', '.pdf', '.docx', '.md']
        
    def load_documents(self, doc_dir: str) -> List[SimpleDocument]:
        """加载文档目录中的所有文档"""
        print(f"开始加载文档目录: {doc_dir}")
        documents = []
        doc_path = Path(doc_dir)
        
        if not doc_path.exists():
            print(f"文档目录不存在: {doc_dir}")
            return documents
            
        # 遍历文档目录
        for file_path in tqdm(doc_path.rglob("*"), desc="加载文档"):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                try:
                    content = self._read_file(file_path)
                    if content.strip():
                        doc = SimpleDocument(
                            text=content,
                            metadata={"filename": file_path.name, "filepath": str(file_path)}
                        )
                        documents.append(doc)
                        print(f"成功加载文档: {file_path.name}")
                except Exception as e:
                    print(f"加载文档失败 {file_path}: {e}")
                    
        print(f"总共加载了 {len(documents)} 个文档")
        return documents
    
    def _read_file(self, file_path: Path) -> str:
        """根据文件类型读取文件内容"""
        suffix = file_path.suffix.lower()
        
        if suffix == '.docx':
            return self._read_docx(file_path)
        elif suffix == '.pdf':
            return self._read_pdf(file_path)
        elif suffix in ['.txt', '.md']:
            return self._read_text_file(file_path)
        else:
            # 尝试作为文本文件读取
            return self._read_text_file(file_path)
    
    def _read_docx(self, file_path: Path) -> str:
        """读取Word文档内容"""
        if not DOCX_AVAILABLE:
            print(f"⚠️ python-docx不可用，跳过: {file_path.name}")
            return ""
        
        try:
            doc = Document(file_path)
            text_content = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            full_text = '\n'.join(text_content)
            print(f"✅ 成功提取DOCX文档: {file_path.name} (长度: {len(full_text)} 字符)")
            return full_text
            
        except Exception as e:
            print(f"❌ 读取docx文件失败 {file_path.name}: {e}")
            return ""
    
    def _read_pdf(self, file_path: Path) -> str:
        """读取PDF文档内容"""
        if not PDF_AVAILABLE:
            print(f"⚠️ PyPDF2不可用，跳过: {file_path.name}")
            return ""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = []
                
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text.strip())
                
                full_text = '\n'.join(text_content)
                print(f"✅ 成功提取PDF文档: {file_path.name} (长度: {len(full_text)} 字符)")
                return full_text
                
        except Exception as e:
            print(f"❌ 读取PDF文件失败 {file_path.name}: {e}")
            return ""
    
    def _read_text_file(self, file_path: Path) -> str:
        """读取文本文件内容，自动检测编码"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'utf-16le', 'utf-16be']
        
        # 如果有chardet，先尝试自动检测编码
        try:
            import chardet
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                detected = chardet.detect(raw_data)
                if detected['encoding'] and detected['confidence'] > 0.7:
                    encodings.insert(0, detected['encoding'])
        except:
            pass
        
        # 尝试不同编码
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    if content.strip():  # 确保不是空内容
                        print(f"✅ 成功读取文本文件: {file_path.name} (编码: {encoding}, 长度: {len(content)} 字符)")
                        return content
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as e:
                print(f"读取文件时出错 {file_path.name} (编码: {encoding}): {e}")
                continue
        
        # 最后尝试忽略错误的方式读取
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
                print(f"⚠️ 使用错误忽略模式读取: {file_path.name} (长度: {len(content)} 字符)")
                return content
        except Exception as e:
            print(f"❌ 文件读取完全失败 {file_path.name}: {e}")
            return ""


class SimpleLLM:
    """简化的LLM类，使用transformers直接实现"""
    
    def __init__(self, model_path: str):
        self.model_path = model_path
        self.model = None
        self.tokenizer = None
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        self.load_model()
    
    def load_model(self):
        """加载模型"""
        if not TRANSFORMERS_AVAILABLE:
            print("⚠️ transformers不可用，无法加载LLM")
            return
            
        try:
            print(f"加载LLM模型: {self.model_path}")
            self.tokenizer = AutoTokenizer.from_pretrained(
                self.model_path,
                trust_remote_code=True
            )
            
            self.model = AutoModelForCausalLM.from_pretrained(
                self.model_path,
                torch_dtype=torch.float16,
                device_map="auto",
                trust_remote_code=True
            )
            print("✅ LLM模型加载成功")
            
        except Exception as e:
            print(f"❌ LLM模型加载失败: {e}")
            self.model = None
            self.tokenizer = None
    
    def generate(self, prompt: str, max_length: int = 2048) -> str:
        """生成回答"""
        if not self.model or not self.tokenizer:
            return "模型未加载，无法生成回答"
            
        try:
            # 构建对话格式
            messages = [
                {"role": "system", "content": "你是一个专业的金融监管制度问答助手。"},
                {"role": "user", "content": prompt}
            ]
            
            # 使用聊天模板
            text = self.tokenizer.apply_chat_template(
                messages,
                tokenize=False,
                add_generation_prompt=True
            )
            
            # 编码输入
            inputs = self.tokenizer.encode(text, return_tensors="pt").to(self.device)
            
            # 生成回答
            with torch.no_grad():
                outputs = self.model.generate(
                    inputs,
                    max_length=min(len(inputs[0]) + max_length, 4096),
                    temperature=0.1,
                    do_sample=True,
                    top_p=0.8,
                    pad_token_id=self.tokenizer.eos_token_id
                )
            
            # 解码回答
            response = self.tokenizer.decode(
                outputs[0][len(inputs[0]):],
                skip_special_tokens=True
            )
            
            return response.strip()
            
        except Exception as e:
            print(f"生成回答时出错: {e}")
            return f"生成回答时出现错误: {e}"


class SimpleEmbedding:
    """简化的嵌入模型类"""
    
    def __init__(self, model_path: str):
        self.model_path = model_path
        self.model = None
        self.load_model()
    
    def load_model(self):
        """加载嵌入模型"""
        try:
            print(f"加载嵌入模型: {self.model_path}")
            self.model = SentenceTransformer(self.model_path)
            print("✅ 嵌入模型加载成功")
        except Exception as e:
            print(f"❌ 嵌入模型加载失败: {e}")
            self.model = None
    
    def get_text_embedding(self, text: str):
        """获取文本嵌入"""
        if not self.model:
            # 返回随机向量作为fallback
            import numpy as np
            return np.random.random(768).tolist()
        
        try:
            embedding = self.model.encode(text)
            return embedding.tolist()
        except Exception as e:
            print(f"嵌入生成失败: {e}")
            import numpy as np
            return np.random.random(768).tolist()
    
    def encode(self, texts: List[str], **kwargs):
        """批量编码文本"""
        if not self.model:
            import numpy as np
            return np.random.random((len(texts), 768))
        
        try:
            return self.model.encode(texts, **kwargs)
        except Exception as e:
            print(f"批量编码失败: {e}")
            import numpy as np
            return np.random.random((len(texts), 768))


class RAGEngine:
    """RAG引擎主类 - 简化版本，专注于向量检索"""
    
    def __init__(self):
        self.config = Config
        self.llm = None
        self.embed_model = None
        self.doc_processor = DocumentProcessor()
        self.vector_db = None
        
        # 初始化模型
        self._init_models()
        
        # 初始化向量数据库
        self._init_vector_db()
        
    def _init_models(self):
        """初始化大模型和嵌入模型"""
        print("正在初始化模型...")
        
        # 初始化嵌入模型
        print("加载嵌入模型...")
        try:
            self.embed_model = SimpleEmbedding(self.config.EMBEDDING_MODEL_PATH)
        except Exception as e:
            print(f"嵌入模型初始化失败: {e}")
            self.embed_model = None
        
        # 初始化大语言模型
        print("加载大语言模型...")
        try:
            self.llm = SimpleLLM(self.config.LLM_MODEL_PATH)
        except Exception as e:
            print(f"大语言模型初始化失败: {e}")
            self.llm = None
        
        # 清理GPU缓存
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
            gc.collect()
    
    def _init_vector_db(self):
        """初始化向量数据库"""
        print("初始化向量数据库...")
        self.vector_db = VectorDatabase(embedding_model=self.embed_model)
    
    def build_index(self, force_rebuild: bool = False):
        """构建或加载向量索引"""
        print("构建向量索引...")
        
        # 尝试从向量数据库加载
        if not force_rebuild and self.vector_db.can_load_existing():
            if self.vector_db.load_from_disk():
                print("从向量数据库加载索引成功")
                stats = self.vector_db.get_statistics()
                print(f"向量数据库统计: {stats}")
                return
        
        # 需要构建新的向量数据库
        print("构建新的向量数据库...")
        
        # 加载文档
        documents = self.doc_processor.load_documents(self.config.DOCUMENTS_DIR)
        if not documents:
            raise ValueError("没有找到可用的文档")
        
        # 转换为向量数据库格式
        vector_db_docs = []
        for doc in documents:
            vector_db_docs.append({
                'text': doc.text,
                'metadata': doc.metadata
            })
        
        # 构建向量数据库
        success = self.vector_db.build_from_documents(vector_db_docs, force_rebuild=True)
        if not success:
            raise ValueError("向量数据库构建失败")
        
        # 显示统计信息
        stats = self.vector_db.get_statistics()
        print(f"向量数据库构建完成: {stats}")
        
        # 清理内存
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
            gc.collect()
    
    def retrieve_documents(self, query: str) -> List[str]:
        """检索相关文档"""
        if not self.vector_db.is_loaded:
            raise ValueError("向量数据库未加载")
            
        print(f"检索查询: {query}")
        
        # 使用向量数据库搜索
        search_results = self.vector_db.search(query, top_k=self.config.TOP_K)
        
        # 提取文档内容
        contexts = []
        for result in search_results:
            contexts.append(result['text'])
            print(f"检索到相关文档片段 (分数: {result['score']:.4f}): {result['text'][:100]}...")
            
        print(f"共检索到 {len(contexts)} 个相关文档片段")
        return contexts
    
    def generate_answer(self, question: str, context: str, question_type: str) -> str:
        """生成答案"""
        print(f"生成答案 - 问题类型: {question_type}")
        
        if question_type == "选择题":
            # 解析选择题的选项
            parts = question.split('\n')
            if len(parts) > 1:
                question_text = parts[0]
                options = '\n'.join(parts[1:])
            else:
                question_text = question
                options = ""
                
            prompt = self.config.CHOICE_PROMPT_TEMPLATE.format(
                context=context,
                question=question_text,
                options=options
            )
        else:
            prompt = self.config.QA_PROMPT_TEMPLATE.format(
                context=context,
                question=question
            )
        
        print(f"生成的提示词长度: {len(prompt)}")
        
        try:
            if self.llm and self.llm.model:
                # 使用LLM生成答案
                answer = self.llm.generate(prompt, max_length=self.config.MAX_TOKENS)
            else:
                # 如果LLM不可用，返回基于检索的简单回答
                answer = f"基于检索到的相关文档，针对问题'{question}'，相关内容如下：\n{context[:500]}..."
            
            print(f"生成的答案: {answer[:200]}...")
            return answer
            
        except Exception as e:
            print(f"生成答案时发生错误: {e}")
            return "抱歉，生成答案时出现错误。"
    
    def answer_question(self, question: str, question_type: str) -> Dict[str, Any]:
        """回答问题的主函数"""
        print(f"\n{'='*50}")
        print(f"开始处理问题: {question[:100]}...")
        print(f"问题类型: {question_type}")
        
        try:
            # 检索相关文档
            contexts = self.retrieve_documents(question)
            
            if not contexts:
                return {
                    "question": question,
                    "answer": "未找到相关文档",
                    "confidence": 0.0,
                    "sources": []
                }
            
            # 合并检索到的文档内容
            combined_context = '\n\n'.join(contexts[:3])  # 使用前3个最相关的文档
            
            # 生成答案
            answer = self.generate_answer(question, combined_context, question_type)
            
            result = {
                "question": question,
                "answer": answer,
                "context_used": combined_context,
                "num_sources": len(contexts)
            }
            
            print(f"问题处理完成")
            return result
            
        except Exception as e:
            print(f"处理问题时发生错误: {e}")
            return {
                "question": question,
                "answer": f"处理问题时发生错误: {e}",
                "error": str(e)
            }
    
    def get_vector_db_stats(self) -> Dict[str, Any]:
        """获取向量数据库统计信息"""
        if self.vector_db:
            return self.vector_db.get_statistics()
        return {"status": "未初始化"}
    
    def rebuild_vector_db(self):
        """重建向量数据库"""
        print("重建向量数据库...")
        if self.vector_db:
            self.vector_db.clear_database()
        self.build_index(force_rebuild=True)
    
    def cleanup(self):
        """清理资源"""
        print("清理系统资源...")
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
        gc.collect()
        print("资源清理完成") 