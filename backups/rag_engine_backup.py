"""
RAG引擎核心实现
包含文档处理、向量检索、答案生成等功能
集成FAISS向量数据库实现高效持久化
更新支持Qwen3-8B模型
"""

import os
import json
import gc
import torch
from pathlib import Path
from typing import List, Dict, Any, Optional
from tqdm import tqdm
import re
import numpy as np
import logging

# 文档处理依赖
try:
    from docx import Document
    DOCX_AVAILABLE = True
    print("✅ python-docx可用")
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx不可用，.docx文件将无法正确处理")

try:
    import PyPDF2
    PDF_AVAILABLE = True
    print("✅ PyPDF2可用")
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️ PyPDF2不可用，.pdf文件将无法正确处理")

# 导入Universal LLM
try:
    from universal_llm import UniversalLLM, create_universal_llm
    UNIVERSAL_LLM_AVAILABLE = True
    print("✅ Universal LLM类导入成功")
except ImportError as e:
    print(f"⚠️ Universal LLM类导入失败: {e}")
    UNIVERSAL_LLM_AVAILABLE = False

# 导入Qwen3专用LLM类（作为备选）
try:
    from qwen3_llm import Qwen3LLM, create_qwen3_llm
    QWEN3_AVAILABLE = True
    print("✅ Qwen3 LLM类导入成功")
except ImportError as e:
    print(f"⚠️ Qwen3 LLM类导入失败: {e}")
    QWEN3_AVAILABLE = False

# 简化导入，避免llama-index的循环导入问题
try:
    # 直接导入transformers和sentence-transformers作为替代
    from transformers import AutoTokenizer, AutoModelForCausalLM
    from sentence_transformers import SentenceTransformer
    TRANSFORMERS_AVAILABLE = True
    print("✅ transformers和sentence-transformers导入成功")
except ImportError as e:
    print(f"⚠️ transformers导入失败: {e}")
    TRANSFORMERS_AVAILABLE = False

# 尝试导入jieba
try:
    import jieba
    JIEBA_AVAILABLE = True
    print("✅ jieba导入成功")
except ImportError:
    print("⚠️ jieba未安装，文本分词功能将受限")
    JIEBA_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
    print("✅ pandas导入成功")
except ImportError:
    print("⚠️ pandas未安装")
    PANDAS_AVAILABLE = False

from config import Config
from vector_db import VectorDatabase

# 增强检索系统导入
try:
    from enhanced_retrieval_system import HybridRetriever, QueryProcessor, RelevanceReranker
    ENHANCED_RETRIEVAL_AVAILABLE = True
    print("✅ 增强检索系统可用")
except ImportError:
    ENHANCED_RETRIEVAL_AVAILABLE = False
    print("⚠️ 增强检索系统不可用")


class SimpleDocument:
    """简化的文档类，替代llama-index的Document"""
    def __init__(self, text: str, metadata: Dict[str, Any] = None):
        self.text = text
        self.metadata = metadata or {}


class DocumentProcessor:
    """文档处理器"""
    
    def __init__(self):
        self.supported_formats = ['.txt', '.pdf', '.docx', '.md']
        
    def load_documents(self, doc_dir: str) -> List[SimpleDocument]:
        """加载文档目录中的所有文档"""
        print(f"开始加载文档目录: {doc_dir}")
        documents = []
        doc_path = Path(doc_dir)
        
        if not doc_path.exists():
            print(f"文档目录不存在: {doc_dir}")
            return documents
            
        # 遍历文档目录
        for file_path in tqdm(doc_path.rglob("*"), desc="加载文档"):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                try:
                    content = self._read_file(file_path)
                    if content.strip():
                        doc = SimpleDocument(
                            text=content,
                            metadata={"filename": file_path.name, "filepath": str(file_path)}
                        )
                        documents.append(doc)
                        print(f"成功加载文档: {file_path.name}")
                except Exception as e:
                    print(f"加载文档失败 {file_path}: {e}")
                    
        print(f"总共加载了 {len(documents)} 个文档")
        return documents
    
    def _read_file(self, file_path: Path) -> str:
        """根据文件类型读取文件内容"""
        suffix = file_path.suffix.lower()
        
        if suffix == '.docx':
            return self._read_docx(file_path)
        elif suffix == '.pdf':
            return self._read_pdf(file_path)
        elif suffix in ['.txt', '.md']:
            return self._read_text_file(file_path)
        else:
            # 尝试作为文本文件读取
            return self._read_text_file(file_path)
    
    def _read_docx(self, file_path: Path) -> str:
        """读取Word文档内容，优化处理方式，特别关注条款完整性和标题保留"""
        if not DOCX_AVAILABLE:
            print(f"⚠️ python-docx不可用，跳过: {file_path.name}")
            return ""
        
        try:
            doc = Document(file_path)
            text_content = []
            current_article = []
            in_article = False
            document_title = ""
            
            # 提取文档标题（通常是第一段或带有标题样式的段落）
            for i, paragraph in enumerate(doc.paragraphs):
                if i == 0 and paragraph.text.strip():
                    document_title = paragraph.text.strip()
                    break
                elif paragraph.style.name.startswith('Title') and paragraph.text.strip():
                    document_title = paragraph.text.strip()
                    break
            
            if document_title:
                text_content.append(f"# {document_title} #")
            
            # 1. 提取标题和正文
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                    
                # 检查是否是条款标题（如"第一条"、"第一章"等）
                if re.match(r'^第[一二三四五六七八九十百]+[章节条]', text) or re.match(r'^第\d+[章节条]', text):
                    # 如果已经在处理条款，先保存当前条款
                    if in_article and current_article:
                        text_content.append('\n'.join(current_article))
                        current_article = []
                    
                    # 开始新的条款
                    in_article = True
                    current_article.append(f"\n## {text} ##\n")
                elif paragraph.style.name.startswith('Heading') or paragraph.style.name.startswith('标题'):
                    # 处理其他标题
                    if in_article:
                        current_article.append(f"\n### {text} ###\n")
                    else:
                        text_content.append(f"\n### {text} ###\n")
                else:
                    if in_article:
                        current_article.append(text)
                    else:
                        # 非条款内容直接添加
                        text_content.append(text)
            
            # 保存最后一个条款
            if current_article:
                text_content.append('\n'.join(current_article))
            
            # 2. 提取表格内容，保持表格结构
            for table in doc.tables:
                table_text = []
                table_text.append("\n【表格开始】")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(" | ".join(row_text))
                table_text.append("【表格结束】\n")
                if len(table_text) > 2:  # 确保表格有内容
                    text_content.append("\n".join(table_text))
            
            # 3. 合并文本，保持段落结构
            full_text = "\n".join(text_content)
            
            # 4. 清理文本
            full_text = self._clean_docx_text(full_text)
            
            print(f"✅ 成功提取DOCX文档: {file_path.name} (长度: {len(full_text)} 字符)")
            return full_text
            
        except Exception as e:
            print(f"❌ 读取docx文件失败 {file_path.name}: {e}")
            return ""
    
    def _clean_docx_text(self, text: str) -> str:
        """清理docx文本，保留标记符号"""
        # 1. 移除多余空白，但保留标题标记
        text = re.sub(r'(?<!#)\s+(?!#)', ' ', text)
        
        # 2. 保留中文标点符号和标记符号
        text = re.sub(r'[^\w\s\u4e00-\u9fff.,，。；：？！""''（）()#|【】]', '', text)
        
        # 3. 处理特殊格式
        text = re.sub(r'\n{3,}', '\n\n', text)  # 移除多余空行
        
        # 4. 确保标题和内容之间有适当的空行
        text = re.sub(r'(###.*?###)\s*', r'\1\n\n', text)
        text = re.sub(r'(##.*?##)\s*', r'\1\n\n', text)
        
        return text.strip()
    
    def _read_pdf(self, file_path: Path) -> str:
        """读取PDF文档内容"""
        if not PDF_AVAILABLE:
            print(f"⚠️ PyPDF2不可用，跳过: {file_path.name}")
            return ""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = []
                
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text.strip())
                
                full_text = '\n'.join(text_content)
                print(f"✅ 成功提取PDF文档: {file_path.name} (长度: {len(full_text)} 字符)")
                return full_text
                
        except Exception as e:
            print(f"❌ 读取PDF文件失败 {file_path.name}: {e}")
            return ""
    
    def _read_text_file(self, file_path: Path) -> str:
        """读取文本文件内容，自动检测编码"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'utf-16le', 'utf-16be']
        
        # 如果有chardet，先尝试自动检测编码
        try:
            import chardet
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                detected = chardet.detect(raw_data)
                if detected['encoding'] and detected['confidence'] > 0.7:
                    encodings.insert(0, detected['encoding'])
        except:
            pass
        
        # 尝试不同编码
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    if content.strip():  # 确保不是空内容
                        print(f"✅ 成功读取文本文件: {file_path.name} (编码: {encoding}, 长度: {len(content)} 字符)")
                        return content
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as e:
                print(f"读取文件时出错 {file_path.name} (编码: {encoding}): {e}")
                continue
        
        # 最后尝试忽略错误的方式读取
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
                print(f"⚠️ 使用错误忽略模式读取: {file_path.name} (长度: {len(content)} 字符)")
                return content
        except Exception as e:
            print(f"❌ 文件读取完全失败 {file_path.name}: {e}")
            return ""

    def _split_document(self, content: str) -> List[str]:
        """优化的文档切片方法，保持条款完整性，减少重合"""
        chunk_size = self.config.CHUNK_SIZE
        chunk_overlap = min(self.config.CHUNK_OVERLAP, chunk_size // 4)  # 限制重叠不超过1/4
        min_chunk_length = getattr(self.config, 'MIN_CHUNK_LENGTH', 50)
        
        if len(content) <= chunk_size:
            return [content] if len(content) >= min_chunk_length else []
        
        chunks = []
        
        # 提取文档标题
        document_title = ""
        title_match = re.search(r'#\s*(.*?)\s*#', content)
        if title_match:
            document_title = title_match.group(1)
        
        # 按条款和标题分割
        sections = re.split(r'(\n##\s*第[一二三四五六七八九十百]+[章节条].*?##\n|\n###.*?###\n)', content)
        
        current_chunk = []
        current_length = 0
        
        # 添加文档标题到每个块
        prefix = f"# {document_title} #\n" if document_title else ""
        prefix_length = len(prefix)
        
        for i in range(0, len(sections)):
            section = sections[i].strip()
            if not section:
                continue
            
            # 检查是否是标题或条款标题
            is_heading = re.match(r'##\s*第[一二三四五六七八九十百]+[章节条].*?##', section) or re.match(r'###.*?###', section)
            
            # 如果是标题，尝试与后面的内容合并
            if is_heading and i + 1 < len(sections):
                section = section + "\n" + sections[i + 1].strip()
                i += 1
            
            # 如果当前块加上这个部分超过限制，先保存当前块
            if current_length + len(section) > chunk_size - prefix_length:
                if current_chunk:
                    full_chunk = prefix + "\n".join(current_chunk)
                    if len(full_chunk) >= min_chunk_length:
                        chunks.append(full_chunk)
                    current_chunk = []
                    current_length = 0
            
            # 如果单个部分超过限制，需要进一步分割
            if len(section) > chunk_size - prefix_length:
                # 先尝试按段落分割
                paragraphs = section.split("\n")
                temp_chunk = []
                temp_length = 0
                
                for para in paragraphs:
                    if temp_length + len(para) <= chunk_size - prefix_length:
                        temp_chunk.append(para)
                        temp_length += len(para) + 1  # +1 for newline
                    else:
                        # 如果段落本身太长，按句子分割
                        if len(para) > chunk_size - prefix_length:
                            sentences = re.split(r'([。！？])', para)
                            for j in range(0, len(sentences), 2):
                                if j + 1 < len(sentences):
                                    sentence = sentences[j] + sentences[j + 1]
                                else:
                                    sentence = sentences[j]
                                    
                                if temp_length + len(sentence) <= chunk_size - prefix_length:
                                    temp_chunk.append(sentence)
                                    temp_length += len(sentence)
                                else:
                                    if temp_chunk:
                                        full_chunk = prefix + "\n".join(temp_chunk)
                                        if len(full_chunk) >= min_chunk_length:
                                            chunks.append(full_chunk)
                                    temp_chunk = [sentence]
                                    temp_length = len(sentence)
                        else:
                            if temp_chunk:
                                full_chunk = prefix + "\n".join(temp_chunk)
                                if len(full_chunk) >= min_chunk_length:
                                    chunks.append(full_chunk)
                            temp_chunk = [para]
                            temp_length = len(para)
                
                if temp_chunk:
                    full_chunk = prefix + "\n".join(temp_chunk)
                    if len(full_chunk) >= min_chunk_length:
                        chunks.append(full_chunk)
            else:
                current_chunk.append(section)
                current_length += len(section) + 1  # +1 for newline
        
        # 保存最后一个块
        if current_chunk:
            full_chunk = prefix + "\n".join(current_chunk)
            if len(full_chunk) >= min_chunk_length:
                chunks.append(full_chunk)
        
        return chunks


class SimpleLLM:
    """兼容性LLM类 - 现在内部使用UniversalLLM或Qwen3LLM"""
    
    def __init__(self, model_path: str):
        self.model_path = model_path
        self.universal_llm = None
        self.qwen3_llm = None
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        self.load_model()
    
    def load_model(self):
        """加载模型 - 优先使用UniversalLLM"""
        if UNIVERSAL_LLM_AVAILABLE:
            print(f"使用UniversalLLM加载模型: {self.model_path}")
            try:
                self.universal_llm = create_universal_llm(self.model_path)
                print("✅ UniversalLLM模型加载成功")
                return
            except Exception as e:
                print(f"❌ UniversalLLM加载失败: {e}")
                self.universal_llm = None
        
        # 如果UniversalLLM不可用，回退到Qwen3LLM
        if QWEN3_AVAILABLE:
            print(f"回退到Qwen3LLM加载模型: {self.model_path}")
            try:
                self.qwen3_llm = create_qwen3_llm(self.model_path)
                print("✅ Qwen3LLM模型加载成功")
                return
            except Exception as e:
                print(f"❌ Qwen3LLM加载失败: {e}")
                self.qwen3_llm = None
        
        # 如果都不可用，回退到原有方式
        if not TRANSFORMERS_AVAILABLE:
            print("⚠️ transformers不可用，无法加载LLM")
            return
            
        try:
            print(f"回退到传统transformers加载: {self.model_path}")
            from transformers import AutoTokenizer, AutoModelForCausalLM
            
            self.tokenizer = AutoTokenizer.from_pretrained(
                self.model_path,
                trust_remote_code=True
            )
            
            self.model = AutoModelForCausalLM.from_pretrained(
                self.model_path,
                torch_dtype=torch.float16,
                device_map="auto",
                trust_remote_code=True
            )
            print("✅ 传统LLM模型加载成功")
            
        except Exception as e:
            print(f"❌ 传统LLM模型加载失败: {e}")
            self.model = None
            self.tokenizer = None
    
    def generate(self, prompt: str, max_length: int = 2048) -> str:
        """生成回答 - 优先使用UniversalLLM"""
        if self.universal_llm and self.universal_llm.loaded:
            # 使用UniversalLLM
            return self.universal_llm.generate(prompt, max_length=max_length)
        
        if self.qwen3_llm and self.qwen3_llm.loaded:
            # 使用Qwen3LLM（不启用thinking）
            return self.qwen3_llm.generate(prompt, max_length=max_length)
        
        # 回退到传统方式
        if not hasattr(self, 'model') or not self.model or not hasattr(self, 'tokenizer') or not self.tokenizer:
            return "模型未加载，无法生成回答"
            
        try:
            # 构建对话格式
            messages = [
                {"role": "system", "content": Config.FINANCIAL_SYSTEM_PROMPT},
                {"role": "user", "content": prompt}
            ]
            
            # 使用聊天模板
            text = self.tokenizer.apply_chat_template(
                messages,
                tokenize=False,
                add_generation_prompt=True
            )
            
            # 编码输入
            inputs = self.tokenizer.encode(text, return_tensors="pt").to(self.device)
            
            # 生成回答 - 使用Qwen2.5配置
            gen_config = Config.QWEN25_GENERATION_CONFIG.copy()
            gen_config["max_new_tokens"] = min(max_length, gen_config["max_new_tokens"])
            
            with torch.no_grad():
                outputs = self.model.generate(
                    inputs,
                    **gen_config
                )
            
            # 解码回答
            response = self.tokenizer.decode(
                outputs[0][len(inputs[0]):],
                skip_special_tokens=True
            )
            
            return response.strip()
            
        except Exception as e:
            print(f"生成回答时出错: {e}")
            return f"生成回答时出现错误: {e}"
    
    def generate_for_financial_qa(self, question: str, context: str, question_type: str) -> str:
        """专门为金融问答生成答案"""
        if self.universal_llm and self.universal_llm.loaded:
            # 使用UniversalLLM的专门方法
            if question_type == "选择题":
                return self.universal_llm.generate_for_choice_question(question, context)
            else:
                return self.universal_llm.generate_for_qa_question(question, context)
        
        if self.qwen3_llm and self.qwen3_llm.loaded:
            # 使用Qwen3的专门方法
            if question_type == "选择题":
                return self.qwen3_llm.generate_for_choice_question(question, context)
            else:
                return self.qwen3_llm.generate_for_qa_question(question, context)
        
        # 回退到传统prompt方式
        if question_type == "选择题":
            # 解析选择题的选项
            parts = question.split('\n')
            if len(parts) > 1:
                question_text = parts[0]
                options = '\n'.join(parts[1:])
            else:
                question_text = question
                options = ""
                
            prompt = Config.CHOICE_PROMPT_TEMPLATE.format(
                context=context,
                question=question_text,
                options=options
            )
        else:
            prompt = Config.QA_PROMPT_TEMPLATE.format(
                context=context,
                question=question
            )
        
        return self.generate(prompt)


class SimpleEmbedding:
    """简化的嵌入模型类"""
    
    def __init__(self, model_path: str):
        self.model_path = model_path
        self.model = None
        self.load_model()
    
    def load_model(self):
        """加载嵌入模型"""
        try:
            print(f"加载嵌入模型: {self.model_path}")
            self.model = SentenceTransformer(self.model_path)
            print("✅ 嵌入模型加载成功")
        except Exception as e:
            print(f"❌ 嵌入模型加载失败: {e}")
            self.model = None
    
    def get_text_embedding(self, text: str):
        """获取文本嵌入"""
        if not self.model:
            # 返回随机向量作为fallback
            return np.random.random(768).tolist()
        
        try:
            embedding = self.model.encode(text)
            return embedding.tolist()
        except Exception as e:
            print(f"嵌入生成失败: {e}")
            return np.random.random(768).tolist()
    
    def encode(self, texts: List[str], **kwargs):
        """批量编码文本"""
        if not self.model:
            return np.random.random((len(texts), 768))
        
        try:
            return self.model.encode(texts, **kwargs)
        except Exception as e:
            print(f"批量编码失败: {e}")
            return np.random.random((len(texts), 768))


class RAGEngine:
    """RAG引擎主类 - 简化版本，专注于向量检索，支持Qwen3"""
    
    def __init__(self, config: Config):
        self.config = config
        self.logger = self._init_logger()
        self.embedding_model = self._init_embedding_model()
        self.llm = self._init_llm()
        self.vector_db = self._init_vector_db()
        # 添加查询缓存
        self.query_cache = {}
        self.cache_size = 100  # 缓存大小限制
        
    def _init_embedding_model(self) -> Any:
        """初始化嵌入模型"""
        try:
            # 设置GPU显存使用比例
            if self.config.USE_GPU and torch.cuda.is_available():
                torch.cuda.set_per_process_memory_fraction(self.config.GPU_MEMORY_FRACTION)
            
            # 加载模型
            model = SentenceTransformer(
                self.config.EMBEDDING_MODEL_PATH,
                device=self.config.MODEL_DEVICE if self.config.USE_GPU else "cpu"
            )
            
            # 设置模型参数
            model.max_seq_length = 512  # 设置最大序列长度
            model.eval()  # 设置为评估模式
            
            self.logger.info(f"嵌入模型 {self.config.EMBEDDING_MODEL_PATH} 初始化成功")
            return model
        
        except Exception as e:
            self.logger.error(f"嵌入模型初始化失败: {str(e)}")
            raise
            
    def _init_llm(self) -> Any:
        """初始化大语言模型"""
        try:
            # 设置GPU显存使用比例
            if self.config.USE_GPU and torch.cuda.is_available():
                torch.cuda.set_per_process_memory_fraction(self.config.GPU_MEMORY_FRACTION)
            
            # 加载模型配置
            model_kwargs = {
                "trust_remote_code": True,
                "device_map": "auto",
                "torch_dtype": getattr(torch, self.config.MODEL_DTYPE),
            }
            
            # 添加量化配置
            if self.config.MODEL_LOAD_IN_8BIT:
                model_kwargs["load_in_8bit"] = True
            elif self.config.MODEL_LOAD_IN_4BIT:
                model_kwargs["load_in_4bit"] = True
            
            # 加载模型
            model = AutoModelForCausalLM.from_pretrained(
                self.config.LLM_MODEL_PATH,
                **model_kwargs
            )
            
            # 加载tokenizer
            tokenizer = AutoTokenizer.from_pretrained(
                self.config.LLM_MODEL_PATH,
                trust_remote_code=True
            )
            
            self.logger.info(f"大语言模型 {self.config.LLM_MODEL_PATH} 初始化成功")
            return {"model": model, "tokenizer": tokenizer}
        
        except Exception as e:
            self.logger.error(f"大语言模型初始化失败: {str(e)}")
            raise
            
    def _init_vector_db(self) -> VectorDatabase:
        """初始化向量数据库"""
        try:
            # 设置GPU显存使用比例
            if self.config.USE_GPU and torch.cuda.is_available():
                torch.cuda.set_per_process_memory_fraction(self.config.GPU_MEMORY_FRACTION)
            
            # 初始化向量数据库
            vector_db = VectorDatabase(
                embedding_model=self.embedding_model
            )
            
            self.logger.info("向量数据库初始化成功")
            return vector_db
        
        except Exception as e:
            self.logger.error(f"向量数据库初始化失败: {str(e)}")
            raise
            
    def _init_logger(self) -> logging.Logger:
        """初始化日志记录器"""
        logger = logging.getLogger("RAGEngine")
        logger.setLevel(logging.INFO)
        handler = logging.StreamHandler()
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        handler.setFormatter(formatter)
        logger.addHandler(handler)
        return logger

    def build_index(self, force_rebuild: bool = False) -> bool:
        """构建知识库索引
        
        Args:
            force_rebuild: 是否强制重建索引
            
        Returns:
            bool: 是否成功构建索引
        """
        try:
            if not force_rebuild and self.vector_db.can_load_existing():
                self.logger.info("使用现有索引")
                return True
                
            self.logger.info("开始构建索引...")
            
            # 加载文档
            doc_processor = DocumentProcessor()
            doc_processor.config = self.config  # 传递配置
            documents = doc_processor.load_documents(self.config.DOCUMENTS_DIR)
            
            if not documents:
                self.logger.error("没有找到可处理的文档")
                return False
            
            # 预处理文档
            processed_docs = self._preprocess_documents(documents)
            
            # 构建索引
            self.vector_db.build_from_documents(processed_docs)
            
            self.logger.info(f"成功构建索引，包含 {len(processed_docs)} 个文档")
            return True
            
        except Exception as e:
            self.logger.error(f"构建索引失败: {e}")
            return False
    
    def _preprocess_documents(self, documents: List[SimpleDocument]) -> List[SimpleDocument]:
        """文档预处理"""
        processed_docs = []
        for doc in documents:
            # 1. 清理文本
            cleaned_text = self._clean_text(doc.text)
            
            # 2. 识别文档结构
            sections = self._identify_sections(cleaned_text)
            
            # 3. 处理每个章节
            for section in sections:
                # 3.1 识别条款
                clauses = self._identify_clauses(section)
                
                # 3.2 处理每个条款
                for clause in clauses:
                    # 3.3 如果条款太长，进行子句切分
                    if len(clause) > self.config.CHUNK_SIZE:
                        sub_clauses = self._split_clause(clause)
                        for sub_clause in sub_clauses:
                            processed_docs.append(SimpleDocument(
                                text=sub_clause,
                                metadata=doc.metadata
                            ))
                    else:
                        processed_docs.append(SimpleDocument(
                            text=clause,
                            metadata=doc.metadata
                        ))
        
        return processed_docs
        
    def _clean_text(self, text: str) -> str:
        """清理文本"""
        # 1. 移除多余空白
        text = re.sub(r'\s+', ' ', text)
        
        # 2. 移除特殊字符
        text = re.sub(r'[^\w\s\u4e00-\u9fff.,，。；：？！""''（）()]', '', text)
        
        # 3. 统一标点符号
        text = text.replace('，', ',').replace('。', '.')
        
        return text.strip()
        
    def _identify_sections(self, content: str) -> List[str]:
        """识别文档章节"""
        sections = []
        current_section = []
        
        for line in content.split('\n'):
            # 识别章节标题
            if re.match(r'^第[一二三四五六七八九十]+[章节]', line):
                if current_section:
                    sections.append('\n'.join(current_section))
                current_section = [line]
            else:
                current_section.append(line)
        
        if current_section:
            sections.append('\n'.join(current_section))
        
        return sections
        
    def _identify_clauses(self, section: str) -> List[str]:
        """识别条款"""
        clauses = []
        current_clause = []
        
        for line in section.split('\n'):
            # 识别条款编号
            if re.match(r'^第\d+条', line):
                if current_clause:
                    clauses.append('\n'.join(current_clause))
                current_clause = [line]
            else:
                current_clause.append(line)
        
        if current_clause:
            clauses.append('\n'.join(current_clause))
        
        return clauses
        
    def _split_clause(self, clause: str) -> List[str]:
        """切分长条款"""
        # 1. 按句号切分
        sentences = re.split(r'[。！？]', clause)
        
        # 2. 合并短句
        chunks = []
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            if current_length + len(sentence) <= self.config.CHUNK_SIZE:
                current_chunk.append(sentence)
                current_length += len(sentence)
            else:
                if current_chunk:
                    chunks.append('。'.join(current_chunk) + '。')
                current_chunk = [sentence]
                current_length = len(sentence)
        
        if current_chunk:
            chunks.append('。'.join(current_chunk) + '。')
        
        return chunks

        def retrieve_documents(self, query: str, top_k: int = 3) -> str:
        """检索相关文档
        
        Args:
            query: 查询文本
            top_k: 返回的文档数量
            
        Returns:
            str: 合并的相关文档文本
        """
        try:
            # 使用基础向量检索
            results = self.vector_db.search(query, top_k=top_k)
            
            if not results:
                return ""
            
            # 合并文档文本
            context = "\n\n".join([doc["text"] for doc in results])
            return context
            
        except Exception as e:
            self.logger.error(f"检索文档失败: {e}")
            return ""
    
    def generate_answer(self, question: str, context: str, question_type: str) -> str:
        """生成答案
        
        Args:
            question: 问题文本
            context: 相关文档文本
            question_type: 问题类型
            
        Returns:
            str: 生成的答案
        """
        try:
            if not self.llm:
                raise ValueError("LLM模型未初始化")
            
            # 构建提示词
            prompt = self._build_prompt(question, context, question_type)
            
            # 生成答案
            answer = self.llm.generate(prompt)
            
            return answer
            
        except Exception as e:
            self.logger.error(f"生成答案失败: {e}")
            raise
    
    def _build_prompt(self, question: str, context: str, question_type: str) -> str:
        """构建提示词
        
        Args:
            question: 问题文本
            context: 相关文档文本
            question_type: 问题类型
            
        Returns:
            str: 构建的提示词
        """
        if question_type == "选择题":
            return f"""请根据以下参考信息回答问题。这是一个选择题，请只给出选项字母（如A、B、C、D）。

参考信息：
{context}

问题：{question}

请只回答选项字母，不要包含其他内容。"""
        else:
            return f"""请根据以下参考信息回答问题。请给出准确、简洁的回答。

参考信息：
{context}

问题：{question}

请直接回答问题，不要包含"根据参考信息"等前缀。"""
    
    def answer_question(self, question: str, question_type: str = "问答题") -> Dict[str, Any]:
        """回答单个问题
        
        Args:
            question: 问题文本
            question_type: 问题类型（选择题/问答题）
            
        Returns:
            Dict[str, Any]: 包含答案和相关信息的字典
        """
        try:
            # 检索相关文档
            context = self.retrieve_documents(question)
            
            if not context:
                return {
                    "answer": "抱歉，没有找到相关的参考信息。",
                    "context_used": "",
                    "num_sources": 0
                }
            
            # 生成答案
            answer = self.generate_answer(question, context, question_type)
            
            # 后处理答案
            processed_answer = self._post_process_answer(answer)
            
            return {
                "answer": processed_answer,
                "context_used": context,
                "num_sources": len(context.split("\n\n"))
            }
            
        except Exception as e:
            self.logger.error(f"回答问题失败: {e}")
            return {
                "answer": f"处理问题时出错: {str(e)}",
                "context_used": "",
                "num_sources": 0,
                "error": str(e)
            }
    
    def _post_process_answer(self, answer: str) -> str:
        """后处理答案"""
        # 1. 清理文本
        answer = self._clean_text(answer)
        
        # 2. 移除重复内容
        answer = self._remove_repetitions(answer)
        
        # 3. 格式化
        answer = self._format_answer(answer)
        
        return answer
        
    def _remove_repetitions(self, text: str) -> str:
        """移除重复内容"""
        sentences = re.split(r'[。！？]', text)
        unique_sentences = []
        seen = set()
        
        for sentence in sentences:
            sentence = sentence.strip()
            if sentence and sentence not in seen:
                unique_sentences.append(sentence)
                seen.add(sentence)
                
        return '。'.join(unique_sentences) + '。'
        
    def _format_answer(self, text: str) -> str:
        """格式化答案"""
        # 1. 分段
        paragraphs = text.split('\n')
        
        # 2. 清理每段
        cleaned_paragraphs = []
        for para in paragraphs:
            para = para.strip()
            if para:
                cleaned_paragraphs.append(para)
                
        # 3. 重新组合
        formatted = '\n\n'.join(cleaned_paragraphs)
        
        return formatted

    def get_vector_db_stats(self) -> Dict[str, Any]:
        """获取向量数据库统计信息"""
        try:
            stats = self.vector_db.get_statistics()
            return {
                "total_documents": stats.get("document_count", 0),
                "total_vectors": stats.get("total_vectors", 0),
                "index_type": stats.get("index_type", "unknown"),
                "status": stats.get("status", "unknown")
            }
        except Exception as e:
            self.logger.error(f"获取向量数据库统计信息失败: {e}")
            return {
                "total_documents": 0,
                "total_vectors": 0,
                "index_type": "unknown",
                "status": "error",
                "error": str(e)
            }
    
    def rebuild_vector_db(self):
        """重建向量数据库"""
        try:
            self.logger.info("开始重建向量数据库...")
            self.build_index(force_rebuild=True)
            self.logger.info("向量数据库重建完成")
        except Exception as e:
            self.logger.error(f"重建向量数据库失败: {e}")
            raise
    
    def cleanup(self):
        """清理资源"""
        try:
            if self.llm and hasattr(self.llm, 'cleanup'):
                self.llm.cleanup()
            if self.embedding_model:
                del self.embedding_model
            torch.cuda.empty_cache()
            gc.collect()
        except Exception as e:
            self.logger.error(f"清理资源失败: {e}")
    
    def search_documents(self, query: str, top_k: int = 3) -> List[Dict[str, Any]]:
        """搜索文档
        
        Args:
            query: 查询文本
            top_k: 返回的文档数量
            
        Returns:
            List[Dict[str, Any]]: 搜索结果列表
        """
        try:
            results = self.vector_db.search(query, top_k=top_k)
            return results
            
        except Exception as e:
            self.logger.error(f"搜索文档失败: {e}")
            return [] 